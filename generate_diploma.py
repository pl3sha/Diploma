# -*- coding: utf-8 -*-
"""
Генератор пояснительной записки к бакалаврской работе.
Тема: веб-система генерации 8-bit персонажей (Stable Diffusion + LoRA).

Запуск: python generate_diploma.py
Результат: diploma_8bit_generator.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

from diploma_sections import (
    CHAPTER_1,
    CHAPTER_2,
    CHAPTER_3,
    CHAPTER_4,
    CONCLUSION,
    ABBREVIATIONS,
    BIBLIOGRAPHY,
)
from diploma_expansion import (
    EXTRA_CH1,
    EXTRA_CH2,
    EXTRA_CH3,
    EXTRA_CH4,
    CH1_APPENDIX_TEXT,
    CH2_APPENDIX_TEXT,
    CH3_APPENDIX_TEXT,
    CH4_APPENDIX_TEXT,
    VOLUME_CH1,
    VOLUME_CH2,
    VOLUME_CH3,
    VOLUME_CH4,
)
from diploma_plain import (
    READER_GUIDE,
    INTRODUCTION_EXTRA,
    CHAPTER_1_INTRO,
    CHAPTER_2_INTRO,
    CHAPTER_3_INTRO,
    CHAPTER_4_INTRO,
    WALKTHROUGH,
    CONCLUSION_EXTRA,
)


INTRODUCTION = """
Современная индустрия независимой разработки видеоигр переживает устойчивый рост: ежегодно на цифровых площадках Steam, itch.io и Google Play появляются десятки тысяч новых проектов, значительная часть которых использует двумерную пиксельную графику. Пиксель-арт сохраняет популярность благодаря узнаваемому художественному языку, сниженным требованиям к анимации и простоте интеграции ресурсов в игровые движки Unity, Godot и GameMaker.

Однако создание качественных спрайтов персонажей остаётся трудоёмким процессом. Для соло-разработчика или небольшой команды без профильного художника этот этап часто становится «узким местом» производственного цикла. За последние годы генеративные модели машинного обучения, в частности Stable Diffusion, стали доступны для локального использования на потребительских видеокартах, что открывает возможность автоматизированной генерации графики по текстовому описанию.

Актуальность данной работы обусловлена сочетанием трёх факторов: сохраняющегося спроса на пиксельную графику в игровой индустрии; доступности open-source инструментов генерации изображений; отсутствия готовых веб-решений, которые объединяют дообучение LoRA на собственном датасете спрайтов и простой пользовательский интерфейс в едином программном комплексе.

Целью работы является разработка веб-системы генерации 8-bit игровых персонажей по текстовому описанию с применением Stable Diffusion v1.5 и дообученного LoRA-адаптера.

Для достижения поставленной цели необходимо решить следующие задачи:
1) провести анализ предметной области и существующих инструментов генерации спрайтов;
2) выбрать и описать средства разработки (Python, FastAPI, React, PyTorch, diffusers, kohya_ss);
3) сформировать датасет пар «изображение — описание», выполнить предобработку и обучить LoRA-адаптер;
4) спроектировать и реализовать клиент-серверное веб-приложение для генерации спрайтов;
5) провести тестирование разработанной системы и оценить качество генерации.

Дополнительно, в рамках исследовательской части, реализована возможность переключения между тремя режимами генерации (базовая модель, публичная LoRA, собственная LoRA) и проведено их наглядное сравнение — это не является основной целью работы, но демонстрирует эффект от дообучения.

Объектом исследования является процесс автоматизированной генерации двумерных спрайтов персонажей по текстовому описанию.

Предметом исследования выступают методы адаптации диффузионной модели Stable Diffusion v1.5 к пиксельному стилю с помощью LoRA и программная архитектура веб-системы.

Структура пояснительной записки: первый раздел посвящён анализу предметной области и постановке задачи. Второй описывает средства разработки. Третий содержит подготовку данных, обучение LoRA и реализацию веб-приложения. Четвёртый — экспериментальную оценку и тестирование. Далее приведены заключение, перечень сокращений, список источников и приложения с листингами программного кода.

Пояснительная записка написана так, чтобы её мог понять читатель без опыта в нейросетях и программировании: специальные термины вводятся постепенно и поясняются простыми словами.
"""


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    r = normal.element.rPr
    if r is not None:
        rFonts = r.rFonts
        if rFonts is not None:
            rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_structural(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_text(doc: Document, text: str) -> None:
    for block in text.strip().split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph(block)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    cap = doc.add_paragraph()
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.add_run(caption).bold = True
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)
    doc.add_paragraph()


def add_code_listing(doc: Document, caption: str, path: Path) -> None:
    cap = doc.add_paragraph()
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.add_run(caption).bold = True
    if path.exists():
        code = path.read_text(encoding="utf-8")
    else:
        code = f"[Файл не найден: {path}]"
    for line in code.splitlines():
        p = doc.add_paragraph(line if line else " ")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for r in p.runs:
            r.font.name = "Courier New"
            r.font.size = Pt(9)


def render_chapter(
    doc: Document,
    sections: list[tuple[str, str]],
    extras: list[tuple[str, str]],
    volume: list[tuple[str, str]],
    appendix: str,
) -> None:
    for title, body in sections:
        add_heading(doc, title)
        add_text(doc, body)
    for title, body in extras:
        add_heading(doc, title)
        add_text(doc, body)
    for title, body in volume:
        add_heading(doc, title)
        add_text(doc, body)
    if appendix.strip():
        add_text(doc, appendix)


def build_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    setup_styles(doc)
    root = Path(__file__).parent

    add_structural(doc, "АННОТАЦИЯ")
    add_text(
        doc,
        """Выпускная квалификационная работа посвящена разработке веб-системы генерации пиксельных (8-bit) игровых персонажей по текстовому описанию с применением нейросетевой модели Stable Diffusion v1.5 и метода эффективного дообучения LoRA.

Объём работы — около 45 страниц пояснительной записки, на которых размещены 5 рисунков и 6 таблиц. При написании работы использовано 20 источников.

Ключевые слова: пиксельная графика, спрайт, Stable Diffusion, LoRA, веб-приложение, FastAPI, React.

Целью работы было создание веб-системы, позволяющей неподготовленному пользователю получать спрайты персонажей в стиле 8-bit по текстовому запросу. Для этого собран датасет из 507 пар «изображение — описание», обучен LoRA-адаптер и реализовано клиент-серверное приложение с REST API.

Разработанное веб-приложение позволяет выполнить полный сценарий генерации без использования командной строки. Дополнительно реализовано переключение между тремя режимами inference для наглядной демонстрации эффекта дообучения.

Результаты могут быть использованы независимыми разработчиками инди-игр для быстрого прототипирования персонажей. Текст пояснительной записки написан доступным языком для широкого круга читателей.""",
    )
    doc.add_page_break()

    add_structural(doc, "СОДЕРЖАНИЕ")
    toc = """ВВЕДЕНИЕ
1 Анализ предметной области и постановка задачи
1.1 Пиксельная графика и спрайты в видеоиграх
1.1.1 Что такое спрайт и пиксель-арт простыми словами
1.1.2 Проблема создания спрайтов для инди-разработчиков
1.2 Обзор существующих решений
1.2.1 Подробный разбор аналога Aseprite
1.2.2 Подробный разбор Automatic1111 WebUI
1.3 Выбор технологии Stable Diffusion и LoRA
1.3.1 Как работает Stable Diffusion — объяснение без формул
1.3.2 Что такое LoRA и зачем она нужна
1.4 Постановка задачи и требования к системе
1.4.1 Пояснение к требованиям простым языком
1.5 Общий принцип работы системы
2 Средства разработки
2.0 Основные понятия для читателя
2.1 Язык Python и фреймворк FastAPI
2.1.1 Клиент-серверная архитектура простыми словами
2.1.2 Библиотека Axios
2.2 Библиотека React и сборщик Vite
2.3 PyTorch и библиотека diffusers
2.3.1 Что такое diffusers и pipeline
2.4 Kohya_ss для обучения LoRA
2.5 Stable Diffusion v1.5 и метод LoRA
2.5.1 Три части Stable Diffusion — простое объяснение
2.6 Словарь терминов работы
3 Подготовка данных, обучение модели и разработка веб-приложения
3.1 Формирование датасета
3.1.1 Откуда взяты изображения и почему это легально
3.2 Предобработка изображений
3.2.1 Пошаговый пример предобработки одного файла
3.2.2 Зачем нужны текстовые описания к каждой картинке
3.3 Обучение LoRA-адаптера
3.3.1 Что такое эпоха, rank и learning rate
3.3.2 Выбор оптимального чекпоинта
3.3.3 График обучения — как его читать
3.4 Архитектура веб-системы
3.4.1 Схема работы системы для неспециалиста
3.5 Реализация серверной части
3.5.1 Описание работы сервера по шагам
3.5.2 REST API — примеры для воспроизведения
3.6 Реализация клиентской части
3.6.1 Подробное описание интерфейса
3.6.2 Как собрать и запустить проект
3.7 Рекомендации по составлению промптов
3.8 Импорт спрайта в игровой движок
4 Тестирование и оценка результатов
4.1 Методика тестирования системы
4.1.1 Критерии оценки качества генерации
4.2 Дополнительное сравнение трёх режимов генерации
4.2.1 Качественный анализ — что видит глаз
4.2.2 Типичные ошибки генерации и как их избежать
4.3 Результаты функционального тестирования
4.3.1 Результаты ручного тестирования
4.3.2 Ограничения системы — честно о минусах
4.4 Сводная таблица: что было сделано
4.5 Рекомендации пользователям и разработчикам
Сквозной пример использования системы
ЗАКЛЮЧЕНИЕ
ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ
СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ
ПРИЛОЖЕНИЕ А Листинг серверной части (backend/main.py)
ПРИЛОЖЕНИЕ Б Листинг клиентской части (frontend/src/App.jsx)
ПРИЛОЖЕНИЕ В Листинг скрипта предобработки (dataset/prepare_dataset.py)
ПРИЛОЖЕНИЕ Г Листинг скрипта оценки (research/evaluate.py)
ПРИЛОЖЕНИЕ Д Листинг скрипта обучения (research/train_lora.py)"""
    add_text(doc, toc)
    doc.add_page_break()

    add_structural(doc, "ВВЕДЕНИЕ")
    add_text(doc, READER_GUIDE)
    add_text(doc, INTRODUCTION)
    add_text(doc, INTRODUCTION_EXTRA)
    doc.add_page_break()

    add_heading(doc, "1 Анализ предметной области и постановка задачи")
    add_text(doc, CHAPTER_1_INTRO)
    render_chapter(doc, CHAPTER_1, EXTRA_CH1, VOLUME_CH1, CH1_APPENDIX_TEXT)
    doc.add_page_break()

    add_heading(doc, "2 Средства разработки")
    add_text(doc, CHAPTER_2_INTRO)
    render_chapter(doc, CHAPTER_2, EXTRA_CH2, VOLUME_CH2, CH2_APPENDIX_TEXT)
    doc.add_page_break()

    add_heading(doc, "3 Подготовка данных, обучение модели и разработка веб-приложения")
    add_text(doc, CHAPTER_3_INTRO)
    render_chapter(doc, CHAPTER_3, EXTRA_CH3, VOLUME_CH3, CH3_APPENDIX_TEXT)

    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Базовая модель", "runwayml/stable-diffusion-v1-5"],
            ["LoRA rank (network_dim)", "32"],
            ["LoRA alpha", "16"],
            ["Число эпох", "15"],
            ["num_repeats", "5"],
            ["Learning rate", "1×10⁻⁴"],
            ["Оптимизатор", "AdamW8bit"],
            ["Точность вычислений", "fp16"],
            ["Размер датасета", "507 пар"],
            ["GPU", "NVIDIA RTX 3070 (8 ГБ VRAM)"],
            ["Активный чекпоинт", "эпоха 14"],
        ],
        "Таблица 3.1 — Параметры обучения LoRA в kohya_ss",
    )

    add_table(
        doc,
        ["Эпоха", "train_loss", "val_loss"],
        [
            ["1", "0,024431", "0,025089"],
            ["5", "0,021045", "0,020199"],
            ["10", "0,019364", "0,023377"],
            ["14", "0,017947", "0,021761"],
            ["15", "0,017332", "0,018106"],
        ],
        "Таблица 3.2 — Фрагмент лога train/val loss (research/loss_log.csv)",
    )
    doc.add_page_break()

    add_heading(doc, "4 Тестирование и оценка результатов")
    add_text(doc, CHAPTER_4_INTRO)
    render_chapter(doc, CHAPTER_4, EXTRA_CH4, VOLUME_CH4, CH4_APPENDIX_TEXT)

    add_table(
        doc,
        ["Критерий", "Base SD v1.5", "Public LoRA", "Custom LoRA"],
        [
            ["Пиксельные контуры", "Размытые", "Частично чёткие", "Чёткие"],
            ["Плоские цвета", "Нет, градиенты", "Частично", "Да"],
            ["Белый фон", "Часто серый", "Обычно белый", "Стабильно белый"],
            ["Соответствие промпту", "Среднее", "Хорошее", "Хорошее"],
            ["Стиль датасета проекта", "Нет", "Частично", "Да"],
        ],
        "Таблица 4.1 — Качественное сравнение трёх режимов (дополнительное исследование)",
    )

    add_table(
        doc,
        ["Число шагов (steps)", "Время на RTX 3070", "Качество"],
        [
            ["10", "~5 с", "Черновое, заметные артефакты"],
            ["25", "~12 с", "Приемлемое для прототипирования"],
            ["50", "~25 с", "Наилучшая детализация"],
        ],
        "Таблица 4.2 — Зависимость времени генерации от числа шагов",
    )

    add_table(
        doc,
        ["Критерий", "Aseprite", "Automatic1111", "Midjourney", "Наша система"],
        [
            ["Автогенерация", "−", "+", "+", "+"],
            ["Специализация под спрайты", "+", "−", "−", "+"],
            ["Сравнение моделей", "−", "−", "−", "+"],
            ["Веб-интерфейс", "−", "−", "+", "+"],
            ["Локальное развёртывание", "+", "+", "−", "+"],
            ["Собственный LoRA", "−", "+", "−", "+"],
        ],
        "Таблица 1.1 — Сравнение аналогов (перенос из раздела 1.2)",
    )
    doc.add_page_break()

    add_heading(doc, "Сквозной пример использования системы")
    add_text(doc, WALKTHROUGH)
    doc.add_page_break()

    add_structural(doc, "ЗАКЛЮЧЕНИЕ")
    add_text(doc, CONCLUSION)
    add_text(doc, CONCLUSION_EXTRA)
    doc.add_page_break()

    add_structural(doc, "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ")
    add_text(doc, ABBREVIATIONS)
    doc.add_page_break()

    add_structural(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    add_text(doc, BIBLIOGRAPHY)
    doc.add_page_break()

    add_structural(doc, "ПРИЛОЖЕНИЕ А")
    add_text(doc, "Листинг серверной части. Данный файл содержит код сервера: загрузку нейросети, переключение LoRA и обработку HTTP-запросов на генерацию.")
    add_code_listing(doc, "Листинг А.1 — backend/main.py", root / "backend" / "main.py")
    doc.add_page_break()

    add_structural(doc, "ПРИЛОЖЕНИЕ Б")
    add_text(doc, "Листинг клиентской части. Данный файл содержит код веб-интерфейса: форму ввода, кнопки выбора модели и отображение результатов.")
    add_code_listing(doc, "Листинг Б.1 — frontend/src/App.jsx", root / "frontend" / "src" / "App.jsx")
    doc.add_page_break()

    add_structural(doc, "ПРИЛОЖЕНИЕ В")
    add_text(doc, "Листинг скрипта предобработки датасета. Программа обрезает спрайты до квадрата, заменяет прозрачный фон на белый и масштабирует до 512×512.")
    add_code_listing(doc, "Листинг В.1 — dataset/prepare_dataset.py", root / "dataset" / "prepare_dataset.py")
    doc.add_page_break()

    add_structural(doc, "ПРИЛОЖЕНИЕ Г")
    add_text(doc, "Листинг вспомогательного скрипта пакетной генерации. Программа использовалась для дополнительного сравнения трёх режимов inference на одном наборе промптов.")
    add_code_listing(doc, "Листинг Г.1 — research/evaluate.py", root / "research" / "evaluate.py")
    doc.add_page_break()

    add_structural(doc, "ПРИЛОЖЕНИЕ Д")
    add_text(doc, "Листинг скрипта обучения LoRA. Программа обучает адаптер с контролем val_loss и сохраняет чекпоинты в формате kohya.")
    add_code_listing(doc, "Листинг Д.1 — research/train_lora.py", root / "research" / "train_lora.py")

    return doc


def main() -> None:
    doc = build_document()
    out = "diploma_8bit_generator.docx"
    doc.save(out)
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    est_pages = total_chars / 1800
    print(f"Saved: {out}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Characters (approx): {total_chars}")
    print(f"Estimated pages (Times 14, 1.5): {est_pages:.0f}")


if __name__ == "__main__":
    main()
